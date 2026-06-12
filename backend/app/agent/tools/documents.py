"""Document generation tools — real PDF / Excel / Word files (Decision 9).

Files land under DOCUMENTS_DIR/<user_id>/ and are recorded in the documents
table through a scoped write transaction, so listing and download inherit the
same per-user RLS as everything else."""

import json
import uuid
from pathlib import Path

from docx import Document as DocxDocument
from openpyxl import Workbook
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.config import DOCUMENTS_DIR
from app.db.pool import execute_scoped
from app.models import User

_EXTENSIONS = {"pdf": ".pdf", "xlsx": ".xlsx", "docx": ".docx"}


def _target_path(user: User, doc_id: uuid.UUID, filename: str, doc_type: str) -> tuple[Path, str]:
    safe = "".join(c for c in filename if c.isalnum() or c in "-_ .").strip() or "document"
    if not safe.endswith(_EXTENSIONS[doc_type]):
        safe += _EXTENSIONS[doc_type]
    user_dir = DOCUMENTS_DIR / user.user_id
    user_dir.mkdir(parents=True, exist_ok=True)
    return user_dir / f"{doc_id}_{safe}", safe


async def _register(user: User, run_id: uuid.UUID, doc_id: uuid.UUID,
                    filename: str, doc_type: str, path: Path) -> str:
    await execute_scoped(
        user,
        """INSERT INTO documents (id, user_id, company_id, run_id, filename,
                                  doc_type, path)
           VALUES ($1, $2, $3, $4, $5, $6, $7)""",
        doc_id, user.user_id, user.company_id, run_id, filename, doc_type, str(path),
    )
    return json.dumps({
        "status": "created",
        "document_id": str(doc_id),
        "filename": filename,
        "download_url": f"/api/documents/{doc_id}/download",
    })


def make_document_tools(user: User, run_id: uuid.UUID):
    async def generate_excel(
        filename: str,
        sheet_name: str,
        headers: list[str],
        rows: list[list[str]],
    ) -> str:
        """Create a real .xlsx file from tabular data and return its download
        URL. Numbers passed as strings are stored as numbers when they parse.

        Args:
            filename: target file name, e.g. "march_energy.xlsx".
            sheet_name: name of the worksheet.
            headers: column header row.
            rows: data rows (lists of cell values as strings).
        """
        doc_id = uuid.uuid4()
        path, safe = _target_path(user, doc_id, filename, "xlsx")
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name[:31] or "Sheet1"
        ws.append(headers)
        for row in rows:
            converted = []
            for cell in row:
                try:
                    converted.append(float(cell))
                except (ValueError, TypeError):
                    converted.append(cell)
            ws.append(converted)
        for col in ws.columns:
            width = max(len(str(c.value or "")) for c in col) + 2
            ws.column_dimensions[col[0].column_letter].width = min(width, 40)
        wb.save(path)
        return await _register(user, run_id, doc_id, safe, "xlsx", path)

    async def generate_pdf(
        filename: str,
        title: str,
        body_text: str,
        table_headers: list[str] | None = None,
        table_rows: list[list[str]] | None = None,
    ) -> str:
        """Create a real PDF report and return its download URL.

        Args:
            filename: target file name, e.g. "march_report.pdf".
            title: document title.
            body_text: report body; blank lines separate paragraphs.
            table_headers: optional table header row.
            table_rows: optional table data rows.
        """
        doc_id = uuid.uuid4()
        path, safe = _target_path(user, doc_id, filename, "pdf")
        styles = getSampleStyleSheet()
        story = [Paragraph(title, styles["Title"]), Spacer(1, 12)]
        for para in body_text.split("\n\n"):
            if para.strip():
                story.append(Paragraph(para.strip().replace("\n", "<br/>"), styles["BodyText"]))
                story.append(Spacer(1, 8))
        if table_headers and table_rows:
            data = [table_headers] + table_rows
            t = Table(data, repeatRows=1)
            t.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1f5132")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f0f6f1")]),
            ]))
            story.append(Spacer(1, 12))
            story.append(t)
        SimpleDocTemplate(str(path), pagesize=A4).build(story)
        return await _register(user, run_id, doc_id, safe, "pdf", path)

    async def generate_word(
        filename: str,
        title: str,
        body_text: str,
        table_headers: list[str] | None = None,
        table_rows: list[list[str]] | None = None,
    ) -> str:
        """Create a real Word (.docx) document and return its download URL.

        Args:
            filename: target file name, e.g. "march_summary.docx".
            title: document title.
            body_text: body; blank lines separate paragraphs.
            table_headers: optional table header row.
            table_rows: optional table data rows.
        """
        doc_id = uuid.uuid4()
        path, safe = _target_path(user, doc_id, filename, "docx")
        doc = DocxDocument()
        doc.add_heading(title, level=1)
        for para in body_text.split("\n\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
        if table_headers and table_rows:
            table = doc.add_table(rows=1, cols=len(table_headers))
            table.style = "Light Grid Accent 1"
            for i, h in enumerate(table_headers):
                table.rows[0].cells[i].text = str(h)
            for row in table_rows:
                cells = table.add_row().cells
                for i, val in enumerate(row[: len(table_headers)]):
                    cells[i].text = str(val)
        doc.save(path)
        return await _register(user, run_id, doc_id, safe, "docx", path)

    return [generate_excel, generate_pdf, generate_word]
